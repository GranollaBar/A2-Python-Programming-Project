import docx
import os
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def buildMemberDocument(username, password, firstname, surname, address, postcode, age, group):
    doc = docx.Document()
    heading = doc.add_heading('Lisburn Racquets Account Details',0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    parag=doc.add_paragraph("")
    parag.add_run("You have successfully been chosen to become a member of Lisburn Racquets Club").bold=True
    parag.add_run("\n\n" + "Your member details are listed below:").bold=True
    parag.add_run("\n\n" + "Username: " + username)
    parag.add_run("\n" + "Password: " + password)
    parag.add_run("\n" + "Firstname: " + firstname)
    parag.add_run("\n" + "Surname: " + surname)
    parag.add_run("\n" + "Address: " + address)
    parag.add_run("\n" + "Postcode: " + postcode)
    parag.add_run("\n" + "Age: " + str(age))
    parag.add_run("\n" + "Group: " + str(group))
    parag.add_run("\n\n" + "Thanks for choosing Lisburn Racquets Club").bold=True
    doc.add_picture('rsz_lisburnraquetsclub.png',width=Inches(3))

    filename = "Member_Account_Details.docx"
    doc.save(filename)
    return open(filename,'rb')
