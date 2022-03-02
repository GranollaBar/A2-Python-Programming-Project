import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import sqlite3

# Builds the most recently completed competition into a word document
def buildcompetitiondocument():
    conn = sqlite3.connect('C:/Users/Josh/pyqt tutorial/AS-Programming-Project/AS Project Frames/_databases_images_doc/Databases/LisburnRacquetsDatabase.db')
    c = conn.cursor()

    c.execute("SELECT * FROM CurrentCompetitionScores")
    row = c.fetchall()

    score1list = []
    score2list = []

    for items in row:
        score1list.append(items[1])
        score2list.append(items[2])

    doc = docx.Document()
    heading = doc.add_heading('Latest Competition Results',0)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    parag=doc.add_paragraph("")
    parag.add_run("These are the results of the most previous " + str(row[0][0]).lower() + ' match with ID: ' + str(row[0][4])).bold=True

    table = doc.add_table(rows=1, cols=2)
    tablerow = table.rows[0].cells
    tablerow[0].text = 'Score 1'
    tablerow[1].text = 'Score 2'

    u = 0

    for values in score1list:
        tablerow = table.add_row().cells
        tablerow[0].text = score1list[u]
        tablerow[1].text = score2list[u]
        u += 1

    table.style = 'Light List'

    parag.add_run("\n\n" + "Final Results").bold=True
    parag.add_run("\n\n" + "Member/Team 1 had a final score of - " + score1list[len(score1list)-1])
    parag.add_run("\n" + "Member/Team 2 had a final score of - " + score2list[len(score2list)-1])

    filename = "C:/Users/Josh/pyqt tutorial/AS-Programming-Project/AS Project Frames/_databases_images_doc/Doc/Current_Competition_Results.docx"
    doc.save(filename)
    return open(filename,'rb')
