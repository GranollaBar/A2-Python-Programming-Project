import docx
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Builds the most recently added member into a word document
def buildMemberDocument(username, password, firstname, surname, telephone, postcode, age, group):
    doc = docx.Document()
    heading = doc.add_heading('Lisburn Racquets Account Details',0)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    parag=doc.add_paragraph("")
    parag.add_run("You have successfully been chosen to become a member of Lisburn Racquets Club").bold=True
    parag.add_run("\n\n" + "Your member details are listed below:").bold=True
    parag.add_run("\n\n" + "Username: " + username)
    parag.add_run("\n" + "Password: " + password)
    parag.add_run("\n" + "Firstname: " + firstname)
    parag.add_run("\n" + "Surname: " + surname)
    parag.add_run("\n" + "Telephone: " + telephone)
    parag.add_run("\n" + "Postcode: " + postcode)
    parag.add_run("\n" + "Age: " + str(age))
    parag.add_run("\n" + "Group: " + str(group))
    parag.add_run("\n\n" + "Thanks for choosing Lisburn Racquets Club").bold=True
    doc.add_picture('C:/Users/Josh/pyqt tutorial/AS-Programming-Project/AS Project Frames/_databases_images_doc/Images/lisburnraquetsclub.png',width=Inches(3))

    filename = "C:/Users/Josh/pyqt tutorial/AS-Programming-Project/AS Project Frames/_databases_images_doc/Doc/Member_Account_Details.docx"
    doc.save(filename)
    return open(filename,'rb')
