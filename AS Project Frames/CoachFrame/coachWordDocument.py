import docx
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Builds the most recently added coach into a word document
def buildCoachDocument(username, password, firstname, surname, gender, DOB, postcode, monday, tuesday, wednesday, thursday, friday, saturday, sunday):
    doc = docx.Document()
    heading = doc.add_heading('Lisburn Racquets Coach Details',0)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    parag=doc.add_paragraph("")
    parag.add_run("You have successfully created your coach account for Lisburn Racquets Club").bold=True
    parag.add_run("\n\n" + "Your details are listed below:").bold=True
    parag.add_run("\n\n" + "Username: " + username)
    parag.add_run("\n" + "Password: " + password)
    parag.add_run("\n" + "Firstname: " + firstname)
    parag.add_run("\n" + "Surname: " + surname)
    parag.add_run("\n" + "Gender: " + gender)
    parag.add_run("\n" + "Date Of Birth: " + DOB)
    parag.add_run("\n" + "Postcode: " + str(postcode))
    parag.add_run("\n" + "Monday Availability: " + str(monday))
    parag.add_run("\n" + "Tuesday Availability: " + str(tuesday))
    parag.add_run("\n" + "Wednesday Availability: " + str(wednesday))
    parag.add_run("\n" + "Thursday Availability: " + str(thursday))
    parag.add_run("\n" + "Friday Availability: " + str(friday))
    parag.add_run("\n" + "Saturday Availability: " + str(saturday))
    parag.add_run("\n" + "Sunday Availability: " + str(sunday))
    parag.add_run("\n\n" + "Thanks for choosing Lisburn Racquets Club").bold=True
    doc.add_picture('C:/Users/Josh/pyqt tutorial/AS-Programming-Project/AS Project Frames/_databases_images_doc/Images/lisburnraquetsclub.png',width=Inches(3))

    filename = "C:/Users/Josh/pyqt tutorial/AS-Programming-Project/AS Project Frames/_databases_images_doc/Doc/Coach_Account_Details.docx"
    doc.save(filename)
    return open(filename,'rb')
